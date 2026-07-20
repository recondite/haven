"""M5 local document ingest: extraction, structuring, draft creation, guards."""
import asyncio
import io

import pytest

from haven import config, docingest, executor, knowledge
from haven import spine as spine_mod
from haven.spine import Spine


def run(coro):
    return asyncio.run(coro)


@pytest.fixture
def env(tmp_path, monkeypatch):
    sb = tmp_path / "SecondBrain"
    (sb / "wiki" / "sources").mkdir(parents=True)
    monkeypatch.setattr(config, "SECONDBRAIN_DIR", sb)
    monkeypatch.setattr(config, "DATA_DIR", tmp_path / "data")
    monkeypatch.setattr(knowledge, "_WIKI_DIR", sb / "wiki")
    monkeypatch.setattr(docingest, "UPLOAD_DIR", tmp_path / "data" / "uploads")
    s = Spine(tmp_path / "spine.sqlite")
    monkeypatch.setattr(spine_mod, "spine", s)
    monkeypatch.setattr(docingest, "spine", s)

    async def fake_call(prompt, model=None, timeout=60.0):
        assert "do not invent" in prompt.lower()
        return "**tl;dr** A test doc.\n\n## Summary\nBudget notes.\n\n## Key facts\n- $2.4M capex"
    monkeypatch.setattr(docingest.runtime, "call", fake_call)
    return s


# ─── extraction ──────────────────────────────────────────
def test_extract_txt_and_md():
    assert "hello" in docingest.extract_text("a.txt", b"hello world")
    assert "# H" in docingest.extract_text("a.md", b"# H\n\nbody")


def test_extract_docx():
    from docx import Document
    doc = Document()
    doc.add_paragraph("First line of the doc.")
    doc.add_paragraph("Second line with a fact: 42 units.")
    buf = io.BytesIO()
    doc.save(buf)
    text = docingest.extract_text("report.docx", buf.getvalue())
    assert "First line" in text and "42 units" in text


def test_unsupported_type_rejected():
    with pytest.raises(docingest.IngestError, match="unsupported"):
        docingest.extract_text("virus.exe", b"MZ...")


# ─── pipeline ────────────────────────────────────────────
def test_ingest_creates_valid_draft_with_provenance(env):
    res = run(docingest.ingest_document(
        "Q3 Budget Memo", "budget.txt", b"Q3 capex is 2.4M, approved by the CFO.", "upload"))
    draft = env.get_draft(res["draft_id"])
    assert draft["kind"] == "wiki"
    assert draft["target"].startswith("wiki/sources/") and draft["target"].endswith(".md")
    # schema-valid (the gate that approve() also runs)
    executor.validate_wiki(draft["payload"], draft["target"])
    assert "## Provenance" in draft["payload"] and "budget.txt" in draft["payload"]
    # raw file preserved append-only under DATA_DIR/uploads
    import json as _json
    ev = _json.loads(draft["evidence"])[0]
    assert (config.DATA_DIR / ev["raw_path"]).exists()


def test_empty_document_rejected(env):
    with pytest.raises(docingest.IngestError, match="no extractable text"):
        run(docingest.ingest_document("x", "empty.txt", b"   ", "upload"))


def test_oversize_rejected(env):
    big = b"x" * (docingest.MAX_BYTES + 1)
    with pytest.raises(docingest.IngestError, match="too large"):
        run(docingest.ingest_document("x", "big.txt", big, "upload"))


def test_duplicate_warning_attached(env):
    # seed an existing page whose title overlaps
    (config.SECONDBRAIN_DIR / "wiki" / "sources" / "q3-budget-memo.md").write_text(
        "---\ntype: source\ncreated: 2026-01-01\nupdated: 2026-01-01\n---\n\n# Q3 Budget Memo\n\nx\n",
        encoding="utf-8")
    res = run(docingest.ingest_document("Q3 Budget Memo notes", "b.txt", b"more budget", "upload"))
    assert res["duplicate_warning"] is True


# ─── M6 google docs ──────────────────────────────────────
def test_extract_file_id():
    assert docingest._extract_file_id(
        "https://docs.google.com/document/d/1AbC_dEf-123/edit") == "1AbC_dEf-123"
    assert docingest._extract_file_id(
        "https://drive.google.com/file/d/9XyZ99/view") == "9XyZ99"
    assert docingest._extract_file_id("https://example.com/nope") is None


def test_gdoc_bad_link(env):
    with pytest.raises(docingest.IngestError, match="Google Doc id"):
        run(docingest.ingest_gdoc("https://example.com/nope"))


def test_gdoc_unauthorized(env, monkeypatch):
    class FakeAuth:
        async def get_drive_service(self):
            return None
    import haven.deps
    monkeypatch.setattr(haven.deps, "gmail_auth", FakeAuth())
    with pytest.raises(docingest.IngestError, match="not authorized"):
        run(docingest.ingest_gdoc("https://docs.google.com/document/d/ABC/edit"))


def test_truncation_marker(env, monkeypatch):
    monkeypatch.setattr(docingest, "_LLM_CHAR_BUDGET", 50)
    res = run(docingest.ingest_document("Long", "long.txt", b"x" * 500, "upload"))
    draft = env.get_draft(res["draft_id"])
    assert "truncated" in draft["payload"]
    assert res["truncated"] is True
